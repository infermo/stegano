import bitstring
import docx
from docx.shared import Pt, RGBColor


def to_binary(file, encoding_type):
    with open(file, "rb") as f:
        text = f.read().decode("UTF-8")
    result = ""
    for symb in text:
        result += bitstring.BitArray(symb.encode(encoding_type)).bin
    return result

#def copy_paragraph_format(source_paragraph, target_paragraph):
#    target_paragraph.style = source_paragraph.style
#    for run in source_paragraph.runs:
#        target_run = target_paragraph.add_run(run.text)
#        target_run.bold = run.bold
#        target_run.italic = run.italic
#        target_run.underline = run.underline
#        target_run.font.color.rgb = run.font.color.rgb if run.font.color else RGBColor(0, 0, 0)
#        target_run.font.size = run.font.size if run.font.size else Pt(30)
#        target_run.font.name = run.font.name if run.font.name else 'Arial'

def hide_message(binary_message, empty_container_name, filled_container_name):
    filled_container = docx.Document()
    empty_container = docx.Document(empty_container_name)
    
    first_paragraph = filled_container.add_paragraph()
    #copy_paragraph_format(empty_container.paragraphs[0], first_paragraph)
    paragraphs = empty_container.paragraphs
    container_text = '\n'.join([p.text for p in paragraphs])
    paragraph = filled_container.add_paragraph()
    for i in range(len(container_text)):
        if container_text[i] == "\n":
            paragraph = filled_container.add_paragraph()
            continue
        if i < len(binary_message) and binary_message[i] == "1":
            run = paragraph.add_run(container_text[i])
            run.font.size = Pt(14.5)
            run.font.name = "Georgia"
        else:
            run = paragraph.add_run(container_text[i])
            run.font.size = Pt(14)
            run.font.name = "Georgia"
    filled_container.save(filled_container_name)

def decode_binary(binary_message, encoding_type):
    if encoding_type == "bo2":  
        return bitstring.BitArray(bin=binary_message).bytes.decode('utf-8')
    else:
        hex_str = hex(int(binary_message, 2))[2:]
        decoded_message = bytes.fromhex(hex_str).decode(encoding_type)
        return decoded_message

if __name__ == "__main__":
    empty_container = "empty_container.docx"
    filled_container = "8.docx"
    encoding_type = "cp866"
    message = "message.txt"
    binary_message = to_binary(message, encoding_type)
    print("Скрываемое сообщение в двоичном виде: " + binary_message)
    hide_message(binary_message, empty_container, filled_container)
    print("Выполнено")
    method = 'font_size'
    print(f"Использованный метод: {method}")
    encodings = ["windows-1251", "koi8-r", "cp866", "bo2"]
    for encoding in encodings:
        if encoding == "bo2":
            continue  
        decoded_message = decode_binary(binary_message, encoding)
        print(f"Декодированное сообщение для {encoding}: {decoded_message}")