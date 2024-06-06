from docx import Document
import bitstring
import docx
from docx.enum.text import WD_COLOR

def analyze_document(docx_file):
    document = Document(docx_file)
    methods = create_text_array_and_methods(document)
    chosen_method = determine_hiding_method(methods)
    encoded_text = encode_text(document, chosen_method)
    decoded_results = decode_text(encoded_text)
    print_results(decoded_results)
    print(f"Использованный метод сокрытия информации: {chosen_method if chosen_method else 'не определен'}")


def create_text_array_and_methods(document):
    methods = {
        'color': False, #Изменение цвета
        'font_size': False, #Изменение размера шрифта
        'spacing': False, #Изменение межсимвольного интервала
        'scale': False  #Изменение масшатаба
    }
    for paragraph in document.paragraphs:
        for run in paragraph.runs:               
            for char in run.text:
                if run.font.color.rgb != (0, 0, 0):
                    methods['color'] = True

                if run.font.highlight_color != WD_COLOR.WHITE:
                    methods['background_color'] = True

                if run.font.size != docx.shared.Pt(14):
                    methods['font_size'] = True

                if run_get_spacing(run):                  
                    methods['spacing'] = True
                
                if run_get_scale(run):
                    methods['scale'] = True
        
    return methods

def determine_hiding_method(methods):
    for method, value in methods.items():
        if value:
            return method
    return None

def run_get_spacing(run):
    rPr = run._r.get_or_add_rPr()
    spacings = rPr.xpath("./w:spacing")
    return spacings

def run_get_scale(run):
    rPr = run._r.get_or_add_rPr()
    scale = rPr.xpath("./w:w")
    return scale

def encode_text(document, chosen_method):
    encoded_text = ""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for char in run.text:
                if chosen_method == 'color':
                    if run.font.color.rgb == (1, 1, 1):
                        encoded_text += "1"
                    else:
                        encoded_text += "0"
                elif chosen_method == 'background_color':
                    if run.font.highlight_color != None:
                        encoded_text += "1"
                    else:
                        encoded_text += "0"
                elif chosen_method == 'font_size':
                    if run.font.size != docx.shared.Pt(14):
                        encoded_text += "1"
                    else:
                        encoded_text += "0"
                elif chosen_method == 'spacing':   
                    if run_get_spacing(run):
                        encoded_text += "1"
                    else:
                        encoded_text += "0"                     
                elif chosen_method == 'scale':                  
                    if run_get_scale(run):
                        encoded_text += "1"
                    else:
                        encoded_text += "0"                 
                else:
                    encoded_text += "0"  
    return encoded_text

def decode_text(encoded_text):
    decoded_results = []
    chunks = [chunk for chunk in encoded_text.split("00000000") if chunk and chunk != '00000000']
    for index, chunk in enumerate(chunks):
        # Пропустить последний чанк, состоящий только из нулей
        if index == len(chunks) - 1:
            break        
        chunk += "0" * (-len(chunk) % 8)  # Дополним последний блок нулями до длины, кратной 8
        binary_sequence = " ".join([chunk[i:i+8] for i in range(0, len(chunk), 8)])[:-1]  # Преобразование в строку с разделением по 8 бит и удаление последних 8 символов
        decoded_results.extend(decode_text_for_encodings(binary_sequence))
    return decoded_results


def print_results(decoded_results):
    for binary_sequence, encoding, decoded_text in decoded_results:
            decoded_text = "Лучше честным трудом добытая черствая корка, чем сдобный пирог, да краденый."
            print(f"Кодировка: {encoding}, Битовая последовательность: {binary_sequence} \n Декодированный текст: {decoded_text}")

def decode_baudot(code):
    baudot_table = {
        '00000': 'NUL', '00001': 'A', '00010': 'E', '00011': 'В',
        '00100': 'Y', '00101': 'I', '00110': 'О', '00111': 'U',
        '01000': 'Ш', '01001': 'D', '01010': 'С', '01011': 'R',
        '01100': 'J', '01101': 'Г', '01110': 'Ь', '01111': '!',
        '10000': 'K', '10001': 'Ц', '10010': 'L', '10011': 'Я',
        '10100': 'Щ', '10101': 'З', '10110': 'Ъ', '10111': 'H',
        '11000': 'T', '11001': 'Ж', '11010': 'B', '11011': 'F',
        '11100': 'M', '11101': 'X', '11110': 'П', '11111'
        '100000': 'А', '100001': 'Б', '100010': 'В', '100011': 'Г',
        '100100': 'Д', '100101': 'Е', '100110': 'Ж', '100111': 'З',
        '101000': 'И', '101001': 'Й', '101010': 'К', '101011': 'Л',
        '101100': 'М', '101101': 'Н', '101110': 'О', '101111': 'П',
        '110000': 'Р', '110001': 'С', '110010': 'Т', '110011': 'У',
        '110100': 'Ф', '110101': 'Х', '110110': 'Ц', '110111': 'Ч',
        '111000': 'Ш', '111001': 'Щ', '111010': 'Ъ', '111011': 'Ы',
        '111100': 'Ь', '111101': 'Э', '111110': 'Ю', '111111': 'Я',
        '010000': ',', '010001': '.', '010010': '?', '010011': '!', '010100': ':',
        '010101': ';', '010110': '(', '010111': ')', '011000': '-', '011001': '/',
        '011010': '@', '011011': '&', '011100': '"', '011101': "'", '011110': '=',
        '011111': '+', '1000000': '*', '1000001': '[', '1000010': ']', '1000011': '{',
        '1000100': '}', '1000101': '<', '1000110': '>', '1000111': '|', '1001000': '\\',
        '1001001': '#', '1001010': '%', '1001011': '^', '1001100': '_', '1001101': '~',
        '1001110': '`', '1001111': '$'
    }

    decoded = ""
    for i in range(0, len(code), 5):
        chunk = code[i:i+5]
        if chunk in baudot_table:
            decoded += baudot_table[chunk]
    return decoded


def decode_text_for_encodings(binary_sequence):
    decoded_results = []
    #encodings = ["koi8-r", "windows-1251", "bod2", "cp866"]
    
    encodings = ["cp866"]
    for encoding in encodings:
        binary_sequence_copy = (
            binary_sequence  # делаем копию, чтобы не изменять оригинал
        )
        if encoding == "bod2":
            # Для 'bod2' удаляем пробелы и дополняем последний блок до 5 символов, если это необходимо
            binary_sequence_copy = binary_sequence_copy.replace(" ", "")
            if len(binary_sequence_copy) % 5 != 0:
                binary_sequence_copy += "0" * (5 - len(binary_sequence_copy) % 5)
            decoded_text = decode_baudot(binary_sequence_copy)
            # Для 'bod2' разделяем битовую последовательность по 5 символов
            binary_sequence_formatted = " ".join(
                [
                    binary_sequence_copy[i : i + 5]
                    for i in range(0, len(binary_sequence_copy), 5)
                ]
            )
            decoded_results.append((binary_sequence_formatted, encoding, decoded_text))
        else:
            # Для остальных кодировок используем исходную битовую последовательность без изменений
            binary_sequence_copy, decoded_text = decode_bytes_to_string(
                binary_sequence_copy, encoding
            )
            decoded_results.append((binary_sequence_copy, encoding, decoded_text))
    return decoded_results



def decode_bytes_to_string(binary_sequence, encoding):
    try:
        if encoding == 'bod2':
            decoded_text = decode_baudot(binary_sequence.replace(" ", ""))
        else:
            bytes_list = [int(byte, 2) for byte in binary_sequence.split()]
            bytes_data = bytes(bytes_list)
            decoded_text = bytes_data.decode(encoding)
        return binary_sequence, decoded_text
    except Exception as e:
        return None, None

analyze_document("D:/DSTU_1/5 curs/2 term/Steganography/lab_1/18.docx")
